"""
Report Generator for Battery System Analysis
- Discovers CSV files in the working directory
- Uses functions from plotv14.py to analyze each CSV and generate plots
- Gathers route photos (morning/evening) if available
- Produces a PDF report (reportlab) and a Word (.docx) report (python-docx)

Usage:
    python report_generator.py

Dependencies:
    pip install reportlab python-docx pillow
"""

import os
import re
import sys
import glob
import math
from datetime import datetime

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
except Exception:
    print("Missing 'reportlab'. Install with: pip install reportlab")

try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    print("Missing 'python-docx'. Install with: pip install python-docx")

from PIL import Image as PILImage
import plotv14
import pandas as pd

# Ordering of shifts
SHIFT_ORDER = {'morning': 0, 'day': 1, 'evening': 2}

# Useful helpers

def find_csv_files(input_dir):
    files = glob.glob(os.path.join(input_dir, "*.csv"))
    return files


def extract_file_metadata(filename):
    # Extract vehicle (prefix before first underscore), shift (morning/evening/day), and date tokens
    base = os.path.basename(filename)
    name = os.path.splitext(base)[0]
    parts = name.split('_')
    vehicle = parts[0] if parts else name
    shift = None
    date_str = None

    # find shift token
    for p in parts:
        if p.lower() in ['morning', 'evening', 'day']:
            shift = p.lower()
            break

    # find a 4-digit year and the preceding tokens to form date
    year_match = re.search(r'(\d{4})', name)
    if year_match:
        year = year_match.group(1)
        # find numbers around year
        tokens = re.findall(r"(\d{1,2})", name)
        # crude but safe: get first two tokens before year if possible
        digits = re.findall(r"\d+", name)
        # try to form date as day_month_year if enough tokens
        if len(digits) >= 3:
            day = digits[-3]
            month = digits[-2]
            date_str = f"{day.zfill(2)}-{month.zfill(2)}-{year}"
        else:
            date_str = year
    else:
        date_str = ''

    return {
        'vehicle': vehicle,
        'shift': shift,
        'date_str': date_str
    }


def find_route_photo(input_dir, vehicle, date_tokens, shift):
    # Look for files that contain vehicle and shift and 'route' (case-insensitive)
    candidates = glob.glob(os.path.join(input_dir, "*"))
    shift_tokens = [shift] if shift else []
    for f in candidates:
        if os.path.isfile(f):
            name = os.path.basename(f).lower()
            if vehicle.lower() in name and (shift and shift in name):
                if 'route' in name or 'map' in name or name.endswith(('.png', '.jpg', '.jpeg')):
                    return f
    return None


def ensure_dir(path):
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)
    return path


def analyze_and_generate(file_path, output_dir, shift_label=None):
    """Run existing analysis functions from plotv14 on the file and save outputs into output_dir.
    Returns a dict with key metrics and list of generated plot file paths."""
    ensure_dir(output_dir)
    # Use plotv14 functions
    print(f"\n📊 Processing file: {os.path.basename(file_path)}")
    df_raw = plotv14.load_and_process_data(file_path)
    if len(df_raw) == 0:
        print(f"No valid data in {file_path}")
        return None

    df_resampled = plotv14.resample_data(df_raw)
    df_resampled, metrics = plotv14.calculate_metrics(df_resampled)

    # Generate plots - pass output_dir so files land here
    print("📈 Generating plots...")
    # Note: plotv14 will name files consistently; to avoid clashes, keep per-shift folder
    plotv14.generate_battery_current_heatmap(df_resampled, output_dir)
    plotv14.generate_power_profile_analysis(df_resampled, output_dir)
    plotv14.generate_battery_current_status_plot(df_raw, output_dir)
    plotv14.generate_voltage_current_plot(df_resampled, output_dir)
    plotv14.generate_permissible_current_plot(df_resampled, output_dir)
    # plotv14.generate_soc_time_plot(df_resampled, output_dir)
    plotv14.generate_ntc_temperatures_plot(df_resampled, output_dir)
    plotv14.generate_soc_vs_voltage_plot(df_resampled, output_dir)
    plotv14.generate_ride_mode_plots(df_resampled, metrics, output_dir)
    plotv14.generate_temperature_distribution_plots(df_resampled, metrics, output_dir)
    plotv14.generate_battery_current_status_plot(df_resampled, output_dir)

    # Collect images
    imgs = []
    for fname in [
        "battery_current_heatmap.png",
        "power_profile_analysis.png",
        "battery_current_status.png",
        "voltage_current_vs_time.png",
        "permissible_vs_battery_current.png",
        # "soc_time.png",
        "ntc_temperatures_vs_time.png",
        "soc_vs_voltage.png",
        "avg_max_current_by_mode.png",
        "avg_max_power_by_mode.png",
        "energy_efficiency_wh_per_km.png",
        "battery_max_temperature_distribution.png",
        "battery_thermal_gradient.png",
    ]:
        p = os.path.join(output_dir, fname)
        if os.path.exists(p):
            imgs.append(p)

    print("✓ Analysis and plot generation complete")
    return {
        'file_path': file_path,
        'output_dir': output_dir,
        'df': df_resampled,
        'metrics': metrics,
        'images': imgs,
        'shift': shift_label
    }


def build_pdf(report_sections, out_pdf_path):
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(out_pdf_path, pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm)
    flow = []

    # Title
    title_style = styles['Title']
    title = Paragraph("BATTERY SYSTEM ANALYSIS REPORT", title_style)
    flow.append(title)
    flow.append(Spacer(1, 6))

    # For each vehicle/date report
    for sec in report_sections:
        # Header
        header = Paragraph(f"<b>Vehicle:</b> {sec['vehicle']} &nbsp;&nbsp; <b>Date:</b> {sec['date']}", styles['Heading2'])
        flow.append(Spacer(1, 6))
        flow.append(header)
        flow.append(Spacer(1, 6))

        # Routes section
        flow.append(Paragraph("<b>Route:</b>", styles['Heading3']))
        for shift in sec['shifts']:
            flow.append(Paragraph(f"<b>{shift['label'].capitalize()} route</b>", styles['Normal']))
            if shift.get('photo') and os.path.exists(shift['photo']):
                img = PILImage.open(shift['photo'])
                iw, ih = img.size
                max_w = 160*mm
                ratio = min(max_w / iw, 1.0)
                img_w_pt = iw * ratio
                img_h_pt = ih * ratio
                rl_img = RLImage(shift['photo'], width=img_w_pt, height=img_h_pt)
                flow.append(rl_img)
                flow.append(Spacer(1, 6))
            else:
                flow.append(Paragraph("No photo available", styles['Normal']))
            flow.append(Spacer(1, 6))

        # Update Electrical Performance Section
        flow.append(Paragraph("<b>1.1 Electrical Performance</b>", styles['Heading3']))
        for shift in sec['shifts']:
            flow.append(Paragraph(f"<b>{shift['label'].capitalize()}</b>", styles['Heading4']))
            m = shift.get('metrics', {})
            vmin, vmax = None, None
            if 'df' in shift and shift['df'] is not None:
                ddf = shift['df']
                if 'Battery_voltage' in ddf.columns:
                    vmin = ddf['Battery_voltage'].min()
                    vmax = ddf['Battery_voltage'].max()
            peak_power = None
            avg_power = None
            try:
                if m and 'df' in shift and shift['df'] is not None:
                    dfd = shift['df']
                    peak_power = dfd['Power_kW'].max()
                    avg_power = dfd['Power_kW'].mean()
            except Exception:
                pass

            perf_lines = []
            if vmin is not None and vmax is not None:
                perf_lines.append(f"Voltage Range: {vmin:.2f}V – {vmax:.2f}V (Δ{(vmax-vmin):.2f}V)")
            if peak_power is not None:
                perf_lines.append(f"Peak Power Output: {peak_power:.2f} kW")
            if avg_power is not None:
                perf_lines.append(f"Average Power: {avg_power:.2f} kW")

            for pl in perf_lines:
                flow.append(Paragraph(pl, styles['Normal']))

        # Update Energy Efficiency Section
        flow.append(Spacer(1, 6))
        flow.append(Paragraph("<b>1.2 Energy Efficiency</b>", styles['Heading3']))
        for shift in sec['shifts']:
            flow.append(Paragraph(f"<b>{shift['label'].capitalize()}</b>", styles['Heading4']))
            metrics_s = shift.get('metrics', {})
            distance = metrics_s.get('distance_km', 0)
            net_energy = metrics_s.get('net_energy_wh', 0)
            discharge = metrics_s.get('cumulative_wh_discharge', 0)
            if distance > 0:
                wh_per_km = net_energy / distance if net_energy != 0 else None
                flow.append(Paragraph(f"Distance Covered: {distance:.2f} km", styles['Normal']))
                flow.append(Paragraph(f"Total Energy Consumed: {discharge:.0f} Wh", styles['Normal']))
                if wh_per_km:
                    flow.append(Paragraph(f"Energy Efficiency: {wh_per_km:.1f} Wh/km", styles['Normal']))

        # Update SOC Analysis Section
        flow.append(Spacer(1, 6))
        flow.append(Paragraph("<b>1.3 SOC Analysis</b>", styles['Heading3']))
        for shift in sec['shifts']:
            flow.append(Paragraph(f"<b>{shift['label'].capitalize()}</b>", styles['Heading4']))
            metrics_s = shift.get('metrics', {})
            soc_initial = metrics_s.get('start_soc', None)
            soc_final = metrics_s.get('end_soc', None)
            if soc_initial is not None and soc_final is not None:
                flow.append(Paragraph(f"Initial SOC: {soc_initial:.1f}%", styles['Normal']))
                flow.append(Paragraph(f"Final SOC: {soc_final:.1f}%", styles['Normal']))
                flow.append(Paragraph(f"SOC Depletion: {(soc_initial - soc_final):.1f}%", styles['Normal']))

        flow.append(Spacer(1, 6))
        # Ride mode analysis table
        flow.append(Paragraph("<b>2. RIDE MODE ANALYSIS</b>", styles['Heading3']))
        # Build table header
        tbl_data = [["Ride Mode", "Avg Current (A)", "Avg Power (W)", "Wh/km", "Usage %"]]
        # aggregate across shifts
        mode_agg = {}
        for s in sec['shifts']:
            metrics_s = s.get('metrics', {})
            if 'mode_stats' in metrics_s:
                # power and current tables
                cur = metrics_s['mode_stats'].get('current') if metrics_s['mode_stats'] else None
                pwr = metrics_s['mode_stats'].get('power') if metrics_s['mode_stats'] else None
                energy = metrics_s['mode_stats'].get('energy') if metrics_s['mode_stats'] else None
                # usage:
                if s.get('df') is not None and 'Ride_Mode' in s['df'].columns:
                    counts = s['df']['Ride_Mode'].value_counts()
                    total = counts.sum()
                    for mode in counts.index:
                        if mode not in mode_agg:
                            mode_agg[mode] = {'count': 0, 'avg_current': [], 'avg_power': [], 'whpkm': []}
                        mode_agg[mode]['count'] += counts[mode]
                        if cur is not None and mode in cur.index:
                            mode_agg[mode]['avg_current'].append(cur.loc[mode, 'Avg_Current_A'])
                        if pwr is not None and mode in pwr.index:
                            mode_agg[mode]['avg_power'].append(pwr.loc[mode, 'Avg_Power_W'])
                        if energy is not None and mode in energy.index:
                            # energy contains Energy_Wh_increment and distance_km
                            whpkm = energy.loc[mode, 'Wh_per_km'] if 'Wh_per_km' in energy.columns else None
                            mode_agg[mode]['whpkm'].append(whpkm)

        total_counts = sum([v['count'] for v in mode_agg.values()]) if mode_agg else 0
        for mode, vals in mode_agg.items():
            avg_curr = (sum(vals['avg_current'])/len(vals['avg_current'])) if vals['avg_current'] else None
            avg_pwr = (sum(vals['avg_power'])/len(vals['avg_power'])) if vals['avg_power'] else None
            avg_wh = (sum([x for x in vals['whpkm'] if x is not None])/len([x for x in vals['whpkm'] if x is not None])) if any(vals['whpkm']) else None
            usage_pct = (vals['count']/total_counts*100) if total_counts>0 else None
            tbl_data.append([mode, f"{avg_curr:.1f}" if avg_curr else "N/A", f"{avg_pwr:.0f}" if avg_pwr else "N/A", f"{avg_wh:.1f}" if avg_wh else "N/A", f"{usage_pct:.1f}%" if usage_pct else "N/A"])

        t = Table(tbl_data, style=[('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('BACKGROUND', (0,0), (-1,0), colors.lightgrey)])
        flow.append(t)
        flow.append(Spacer(1, 12))

        # Thermal section
        flow.append(Paragraph("<b>3. THERMAL PERFORMANCE</b>", styles['Heading3']))
        # thermal table
        thermal_rows = [["Sensor", "Max Temp (°C)", "Avg Temp (°C)", "ΔT from ambient"]]
        # Ambient is not provided; show max and avg
        sensors = ['Ntc_Mos', 'Ntc_Com', 'Ntc_1', 'Ntc_2', 'Ntc_3', 'Ntc_4']
        for s in sensors:
            maxv = None; avgv = None
            for sh in sec['shifts']:
                metrics_sh = sh.get('metrics', {})
                if f'{s}_max' in metrics_sh:
                    maxv = metrics_sh.get(f'{s}_max')
                if f'{s}_avg' in metrics_sh:
                    avgv = metrics_sh.get(f'{s}_avg')
            if maxv is not None:
                thermal_rows.append([s, f"{maxv:.1f}", f"{avgv:.1f}" if avgv else "N/A", "N/A"])
        if len(thermal_rows) > 1:
            tt = Table(thermal_rows, style=[('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('BACKGROUND', (0,0), (-1,0), colors.lightgrey)])
            flow.append(tt)
            flow.append(Spacer(1, 12))

        # Charge/discharge distribution
        flow.append(Paragraph("<b>4. CHARGE/DISCHARGE PROFILE</b>", styles['Heading3']))
        # Use first shift stats if available
        status_rows = []
        for sh in sec['shifts']:
            metrics_sh = sh.get('metrics', {})
            if metrics_sh and 'status_distribution' in metrics_sh:
                sd = metrics_sh['status_distribution']
                status_rows.append([sh['label'], f"Discharging (0): {sd[0]['percentage']:.1f}%" if 0 in sd else "N/A", f"Charging (1): {sd[1]['percentage']:.1f}%" if 1 in sd else "N/A", f"Idle (2): {sd[2]['percentage']:.1f}%" if 2 in sd else "N/A"])
        if status_rows:
            st = Table([['Shift', 'Discharging', 'Charging/Regen', 'Idle']] + status_rows, style=[('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('BACKGROUND', (0,0), (-1,0), colors.lightgrey)])
            flow.append(st)
            flow.append(Spacer(1, 12))

        # Add ride-mode plots grouped by category and shown per shift (morning/day/evening)
        flow.append(Paragraph("<b>Ride Mode Plots</b>", styles['Heading4']))
        # Define categories and matching filename tokens (lowercase)
        categories = [
            ("a. Current Heatmap", ["battery_current_heatmap"]),
            ("b. Power Profile", ["power_profile_analysis"]),
            ("c. Voltage & Current", ["voltage_current_vs_time"]),
            # ("d. Current Status / Permissible", ["battery_current_status", "permissible_vs_battery_current"]),
            # ("e. SOC & SOC-Voltage", ["soc_time", "soc_vs_voltage"]),
            ("d. Current Status / Permissible", ["battery_current_status", "permissible_vs_battery_current"]),
            ("e. SOC & SOC-Voltage", ["soc_vs_voltage"]),
            ("f. Temperature Sensors", ["ntc_temperatures_vs_time", "battery_max_temperature_distribution", "battery_thermal_gradient"]),
            ("g. Mode Summary", ["avg_max_current_by_mode", "avg_max_power_by_mode", "energy_efficiency_wh_per_km"]),
            ("h. Permissible Current", ["permissible_vs_battery_current"]),
        ]

        # Roman numerals for shifts
        romans = ["i.", "ii.", "iii."]

        # Ensure shifts are in the expected order (morning, day, evening)
        shifts_ordered = sec['shifts']

        for idx_cat, (cat_title, tokens) in enumerate(categories, start=1):
            flow.append(Paragraph(f"<b>{cat_title}</b>", styles['Normal']))
            for idx_sh, sh in enumerate(shifts_ordered):
                shift_label = sh.get('label', 'unknown')
                roman = romans[idx_sh] if idx_sh < len(romans) else f"{idx_sh+1}."
                flow.append(Paragraph(f"{roman} {shift_label.capitalize()}", styles['Normal']))

                # Try to find matching image for this category in this shift
                found = False
                for token in tokens:
                    for img in sh.get('images', []):
                        if token in os.path.basename(img).lower():
                            if os.path.exists(img):
                                pil = PILImage.open(img)
                                iw, ih = pil.size
                                max_w = 160*mm
                                ratio = min(max_w / iw, 1.0)
                                rl_img = RLImage(img, width=iw*ratio, height=ih*ratio)
                                flow.append(rl_img)
                                flow.append(Spacer(1, 6))
                                found = True
                                break
                    if found:
                        break
                if not found:
                    flow.append(Paragraph("No plot available for this shift and category.", styles['Normal']))
                flow.append(Spacer(1, 6))

            flow.append(Spacer(1, 12))

        flow.append(Spacer(1, 18))

    doc.build(flow)
    print(f"Saved PDF report: {out_pdf_path}")

def build_word(report_sections, out_docx_path):
    doc = Document()
    doc.add_heading('BATTERY SYSTEM ANALYSIS REPORT', level=1)

    romans = ["i.", "ii.", "iii."]

    def safe_whpkm(mode_stats, mode):
        energy = mode_stats.get('energy')
        if energy is None:
            return None
        # DataFrame case
        if hasattr(energy, "index"):
            if mode in energy.index and 'Wh_per_km' in energy.columns:
                return energy.loc[mode, 'Wh_per_km']
        # Dict case
        if isinstance(energy, dict):
            return energy.get(mode, {}).get('Wh_per_km')
        return None

    for sec in report_sections:
        # ---------------- HEADER ----------------
        doc.add_heading(f"Vehicle: {sec['vehicle']}    Date: {sec['date']}", level=2)

        # ---------------- ROUTE ----------------
        doc.add_heading('Route', level=3)
        for sh in sec['shifts']:
            doc.add_paragraph(f"{sh['label'].capitalize()} route", style='List Bullet')
            if sh.get('photo') and os.path.exists(sh['photo']):
                try:
                    doc.add_picture(sh['photo'], width=Inches(6))
                except Exception:
                    doc.add_paragraph("Could not embed photo")
            else:
                doc.add_paragraph("No photo available")

        # ---------------- 1.1 ELECTRICAL PERFORMANCE ----------------
        doc.add_heading('1.1 Electrical Performance', level=3)
        for sh in sec['shifts']:
            doc.add_heading(sh['label'].capitalize(), level=4)
            df = sh.get('df')
            if df is not None:
                if 'Battery_voltage' in df.columns:
                    vmin = df['Battery_voltage'].min()
                    vmax = df['Battery_voltage'].max()
                    doc.add_paragraph(
                        f"Voltage Range: {vmin:.2f}V – {vmax:.2f}V (Δ{(vmax-vmin):.2f}V)"
                    )
                if 'Power_kW' in df.columns:
                    doc.add_paragraph(f"Peak Power Output: {df['Power_kW'].max():.2f} kW")
                    doc.add_paragraph(f"Average Power: {df['Power_kW'].mean():.2f} kW")

        # ---------------- 1.2 ENERGY EFFICIENCY ----------------
        doc.add_heading('1.2 Energy Efficiency', level=3)
        for sh in sec['shifts']:
            doc.add_heading(sh['label'].capitalize(), level=4)
            m = sh.get('metrics', {})
            if m.get('distance_km', 0) > 0:
                doc.add_paragraph(f"Distance Covered: {m['distance_km']:.2f} km")
                doc.add_paragraph(
                    f"Total Energy Consumed: {m.get('cumulative_wh_discharge', 0):.0f} Wh"
                )
                if m.get('net_energy_wh'):
                    doc.add_paragraph(
                        f"Energy Efficiency: {(m['net_energy_wh']/m['distance_km']):.1f} Wh/km"
                    )

        # ---------------- 1.3 SOC ANALYSIS ----------------
        doc.add_heading('1.3 SOC Analysis', level=3)
        for sh in sec['shifts']:
            doc.add_heading(sh['label'].capitalize(), level=4)
            m = sh.get('metrics', {})
            if m.get('start_soc') is not None and m.get('end_soc') is not None:
                doc.add_paragraph(f"Initial SOC: {m['start_soc']:.1f}%")
                doc.add_paragraph(f"Final SOC: {m['end_soc']:.1f}%")
                doc.add_paragraph(
                    f"SOC Depletion: {(m['start_soc'] - m['end_soc']):.1f}%"
                )

        # ---------------- 2. RIDE MODE ANALYSIS ----------------
        doc.add_heading('2. RIDE MODE ANALYSIS', level=3)

        rm_table = doc.add_table(rows=1, cols=5)
        hdr = rm_table.rows[0].cells
        hdr[0].text = "Ride Mode"
        hdr[1].text = "Avg Current (A)"
        hdr[2].text = "Avg Power (W)"
        hdr[3].text = "Wh/km"
        hdr[4].text = "Usage %"

        mode_agg = {}

        for sh in sec['shifts']:
            df = sh.get('df')
            m = sh.get('metrics', {})
            if df is not None and 'Ride_Mode' in df.columns and 'mode_stats' in m:
                counts = df['Ride_Mode'].value_counts()
                ms = m['mode_stats']

                for mode, cnt in counts.items():
                    if mode not in mode_agg:
                        mode_agg[mode] = {'count': 0, 'curr': [], 'pwr': [], 'wh': []}
                    mode_agg[mode]['count'] += cnt

                    if 'current' in ms and hasattr(ms['current'], 'index'):
                        if mode in ms['current'].index:
                            mode_agg[mode]['curr'].append(
                                ms['current'].loc[mode, 'Avg_Current_A']
                            )

                    if 'power' in ms and hasattr(ms['power'], 'index'):
                        if mode in ms['power'].index:
                            mode_agg[mode]['pwr'].append(
                                ms['power'].loc[mode, 'Avg_Power_W']
                            )

                    whpkm = safe_whpkm(ms, mode)
                    if whpkm is not None:
                        mode_agg[mode]['wh'].append(whpkm)

        total_counts = sum(v['count'] for v in mode_agg.values())

        for mode, v in mode_agg.items():
            r = rm_table.add_row().cells
            r[0].text = str(mode)
            r[1].text = f"{sum(v['curr'])/len(v['curr']):.1f}" if v['curr'] else "N/A"
            r[2].text = f"{sum(v['pwr'])/len(v['pwr']):.0f}" if v['pwr'] else "N/A"
            r[3].text = f"{sum(v['wh'])/len(v['wh']):.1f}" if v['wh'] else "N/A"
            r[4].text = (
                f"{(v['count']/total_counts*100):.1f}%" if total_counts else "N/A"
            )

        # ---------------- 3. THERMAL PERFORMANCE ----------------
        doc.add_heading('3. THERMAL PERFORMANCE', level=3)

        th_table = doc.add_table(rows=1, cols=4)
        th = th_table.rows[0].cells
        th[0].text = "Sensor"
        th[1].text = "Max Temp (°C)"
        th[2].text = "Avg Temp (°C)"
        th[3].text = "ΔT from Ambient"

        sensors = ['Ntc_Mos', 'Ntc_Com', 'Ntc_1', 'Ntc_2', 'Ntc_3', 'Ntc_4']

        for s in sensors:
            maxv, avgv = None, None
            for sh in sec['shifts']:
                m = sh.get('metrics', {})
                if f'{s}_max' in m:
                    maxv = m[f'{s}_max']
                if f'{s}_avg' in m:
                    avgv = m[f'{s}_avg']
            if maxv is not None:
                r = th_table.add_row().cells
                r[0].text = s
                r[1].text = f"{maxv:.1f}"
                r[2].text = f"{avgv:.1f}" if avgv is not None else "N/A"
                r[3].text = "N/A"

        # ---------------- 4. CHARGE / DISCHARGE PROFILE ----------------
        doc.add_heading('4. CHARGE/DISCHARGE PROFILE', level=3)

        sd_table = doc.add_table(rows=1, cols=4)
        sd = sd_table.rows[0].cells
        sd[0].text = "Shift"
        sd[1].text = "Discharging"
        sd[2].text = "Charging/Regen"
        sd[3].text = "Idle"

        for sh in sec['shifts']:
            dist = sh.get('metrics', {}).get('status_distribution')
            if dist:
                r = sd_table.add_row().cells
                r[0].text = sh['label']
                r[1].text = f"{dist[0]['percentage']:.1f}%" if 0 in dist else "N/A"
                r[2].text = f"{dist[1]['percentage']:.1f}%" if 1 in dist else "N/A"
                r[3].text = f"{dist[2]['percentage']:.1f}%" if 2 in dist else "N/A"

        # ---------------- RIDE MODE PLOTS ----------------
        doc.add_heading('Ride Mode Plots', level=4)

        categories = [
            ("a. Current Heatmap", ["battery_current_heatmap"]),
            ("b. Power Profile", ["power_profile_analysis"]),
            ("c. Voltage & Current", ["voltage_current_vs_time"]),
            ("d. Current Status / Permissible", ["battery_current_status", "permissible_vs_battery_current"]),
            ("e. SOC & SOC-Voltage", ["soc_vs_voltage"]),
            ("f. Temperature Sensors", ["ntc_temperatures_vs_time", "battery_max_temperature_distribution", "battery_thermal_gradient"]),
            ("g. Mode Summary", ["avg_max_current_by_mode", "avg_max_power_by_mode", "energy_efficiency_wh_per_km"]),
            ("h. Permissible Current", ["permissible_vs_battery_current"]),
        ]

        for cat, tokens in categories:
            doc.add_paragraph(cat)
            for i, sh in enumerate(sec['shifts']):
                doc.add_paragraph(f"{romans[i]} {sh['label'].capitalize()}")
                found = False
                for token in tokens:
                    for img in sh.get('images', []):
                        if token in os.path.basename(img).lower():
                            try:
                                doc.add_picture(img, width=Inches(6))
                                found = True
                                break
                            except Exception:
                                pass
                    if found:
                        break
                if not found:
                    doc.add_paragraph("No plot available for this shift and category")

        doc.add_page_break()

    doc.save(out_docx_path)
    print(f"Saved Word report: {out_docx_path}")


# ==============================================



def run(input_dir='.', output_root='reports', file_list=None):
    """
    Main driver function.
    :param input_dir: Directory to search for CSVs if file_list is None.
    :param output_root: Directory to save reports.
    :param file_list: Optional list of specific CSV paths to process.
    """
    if file_list:
        csvs = file_list
    else:
        csvs = find_csv_files(input_dir)
        
    if not csvs:
        print('No CSV files found.')
        return

    # Group by vehicle+date
    groups = {}
    for c in csvs:
        meta = extract_file_metadata(c)
        key = (meta['vehicle'], meta['date_str'])
        if key not in groups:
            groups[key] = {'vehicle': meta['vehicle'], 'date': meta['date_str'], 'files': []}
        groups[key]['files'].append({'path': c, 'shift': meta['shift']})

    report_sections = []
    for key, val in groups.items():
        vehicle, date = key
        sec = {'vehicle': vehicle, 'date': date, 'shifts': []}
        # sort files by shift order
        val['files'].sort(key=lambda x: SHIFT_ORDER.get(x.get('shift'), 99))
        for f in val['files']:
            # Determine output folder
            shift_folder = os.path.join(output_root, f"{vehicle}_{date}", f.get('shift') or 'unknown')
            ensure_dir(shift_folder)
            
            # Analyze
            analysis = analyze_and_generate(f['path'], shift_folder, shift_label=f.get('shift') or 'unknown')
            
            # Find photo - Look in the *same directory* as the CSV file first
            csv_dir = os.path.dirname(f['path'])
            photo = find_route_photo(csv_dir, vehicle, date, f.get('shift'))
            
            shift_dict = {
                'label': f.get('shift') or 'unknown',
                'metrics': analysis['metrics'] if analysis else {},
                'images': analysis['images'] if analysis else [],
                'photo': photo,
                'df': analysis['df'] if analysis else None
            }
            sec['shifts'].append(shift_dict)
        report_sections.append(sec)

    # Build reports
    print("\n📝 Building PDF report...")
    ensure_dir(output_root)
    
    # Generate dynamic filename based on content
    vehicles = set(s['vehicle'] for s in report_sections)
    dates = set(s['date'] for s in report_sections)
    
    base_name = "battery_report"
    
    if len(vehicles) == 1 and len(dates) == 1:
        # Single vehicle, single date -> "MULE1_1-7-2026_Report"
        v = list(vehicles)[0]
        d = list(dates)[0]
        base_name = f"{v}_{d}_Report"
    elif len(dates) == 1:
        # Multiple vehicles, single date -> "Combined_1-7-2026_Report"
        d = list(dates)[0]
        base_name = f"Combined_{d}_Report"
    elif len(vehicles) == 1:
        # Single vehicle, multiple dates -> "MULE1_MultiDate_Report"
        v = list(vehicles)[0]
        base_name = f"{v}_MultiDate_Report"
    else:
        # Mixed -> "Battery_Report_TIMESTAMP"
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_name = f"battery_report_{timestamp}"
        
    pdf_path = os.path.join(output_root, f"{base_name}.pdf")
    docx_path = os.path.join(output_root, f"{base_name}.docx")

    build_pdf(report_sections, pdf_path)
    print(f"✓ PDF report saved: {pdf_path}")
    
    print("📝 Building Word report...")
    build_word(report_sections, docx_path)
    print(f"✓ Word report saved: {docx_path}")
    print("\n🎉 Report generation complete!")


if __name__ == "__main__":
    # Automatically set input and output directories to the current working directory
    input_dir = os.getcwd()
    output_root = os.path.join(os.getcwd(), 'reports')
    ensure_dir(output_root)  # Ensure the output directory exists
    
    # Call the run function directly with the updated paths
    run(input_dir=input_dir, output_root=output_root)
